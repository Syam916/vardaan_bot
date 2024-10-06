import streamlit as st
import os
from langchain_groq import ChatGroq  # LLM to handle queries
from langchain_openai import OpenAIEmbeddings  # OpenAI embeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter  # To split large text into smaller chunks
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate  # Prompt template for querying the LLM
from langchain.chains import create_retrieval_chain  # For retrieving relevant documents based on queries
from langchain_community.vectorstores import FAISS  # Vector store for efficient document search
from langchain.schema import Document  # Document schema for handling text data
from dotenv import load_dotenv  # For loading environment variables
import PyPDF2  # For handling PDF extraction
import docx  # For handling DOCX extraction

# Load environment variables like API keys
load_dotenv()

# Load the OpenAI and GROQ API keys
os.environ['OPENAI_API_KEY'] = os.getenv("OPENAI_API_KEY")
groq_api_key = os.getenv('GROQ_API_KEY')

# Set Streamlit page configuration
st.set_page_config(page_title="Vignan Pdf Bot", page_icon="📂", layout="wide")

# Initialize the Large Language Model (GROQ)
llm = ChatGroq(groq_api_key=groq_api_key, model_name="Llama3-8b-8192")

# Define custom prompt template for LLM queries
prompt = ChatPromptTemplate.from_template(
    """
    Answer the questions based on the provided context only.
    Please provide the most accurate response based on the question.
    <context>
    {context}
    <context>
    Questions: {input}
    """
)

# -----------------------------------
# Step 1: Initialize Session States
# This function initializes all the required session states to store data like
# conversation history, document vector store, embeddings, etc.
# -----------------------------------
def initialize_session_states():
    if "conversation" not in st.session_state:
        st.session_state.conversation = []  # Stores chat history
    if "documents_uploaded" not in st.session_state:
        st.session_state.documents_uploaded = False  # Flag for document upload status
    if "query_ready" not in st.session_state:
        st.session_state.query_ready = False  # Flag indicating if query can be processed
    if "vector_store" not in st.session_state:
        st.session_state.vector_store = None  # To store the document embeddings
    if "response" not in st.session_state:
        st.session_state.response = None  # Stores the LLM's response
    if "input_key" not in st.session_state:
        st.session_state.input_key = ""  # For handling input field reset

# -----------------------------------
# Step 2: Extract Text from PDF
# This helper function extracts text content from a PDF file
# -----------------------------------
def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
    return text

# -----------------------------------
# Step 3: Extract Text from DOCX
# This helper function extracts text content from a DOCX file
# -----------------------------------
def extract_text_from_docx(file):
    doc = docx.Document(file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)

# -----------------------------------
# Step 4: Process and Embed Documents
# This function converts documents into embeddings and stores them in the vector store
# It also splits documents into smaller chunks for efficient retrieval during queries
# -----------------------------------
def vector_embedding(docs):
    st.session_state.embeddings = OpenAIEmbeddings()
    st.session_state.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

    # Wrap text in Document objects
    document_objects = [Document(page_content=doc) for doc in docs]

    # Split the documents into chunks
    st.session_state.final_documents = st.session_state.text_splitter.split_documents(document_objects)

    # Create vector store and store it in session_state
    st.session_state.vector_store = FAISS.from_documents(st.session_state.final_documents, st.session_state.embeddings)
    st.success("Documents processed successfully! Ready for querying.")
    st.session_state.query_ready = True  # Querying can start now

# -----------------------------------
# Step 5: Upload and Process Documents
# This function handles the document upload UI, processes uploaded files (PDF, DOCX, TXT)
# and stores them in a vector database after extracting text and creating embeddings.
# -----------------------------------
def handle_document_upload():
    st.header("📂 Upload Documents")
    
    # File uploader for uploading files
    uploaded_files = st.file_uploader(
        "Drag and drop files here (Limit 200MB per file) • PDF, DOCX, TXT", 
        type=["pdf", "docx", "txt"], 
        accept_multiple_files=True, 
        key="file_uploader_1"
    )
    
    if uploaded_files:
        st.session_state.documents_uploaded = True
        st.write(f"{len(uploaded_files)} document(s) uploaded.")
        
        # Proceed button for document embedding
        if st.button("Proceed"):
            with st.spinner("Processing documents..."):
                documents = []
                for uploaded_file in uploaded_files:
                    # Process uploaded files based on their type
                    if uploaded_file.type == "text/plain":
                        text = uploaded_file.read().decode("utf-8")
                        documents.append(text)
                    elif uploaded_file.type == "application/pdf":
                        text = extract_text_from_pdf(uploaded_file)
                        documents.append(text)
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        text = extract_text_from_docx(uploaded_file)
                        documents.append(text)

                # Check if documents have content
                if len(documents) > 0:
                    # Embed documents and create vector store
                    vector_embedding(documents)
                else:
                    st.error("No text extracted from the uploaded documents.")
    else:
        st.write("Upload documents to proceed.")

# -----------------------------------
# Step 6: Handle User Query
# This function manages the user interface for submitting queries,
# retrieves relevant document chunks from the vector store, and displays AI's response.
# -----------------------------------
def handle_query():
    st.header("🤖 Ask Questions")

    if st.session_state.query_ready:
        st.write("You can now ask questions related to the uploaded documents.")
        
        # Chat interface: User input
        user_query = st.text_input("Ask a question", key="input_key", placeholder="Your question...")

        # Handle query submission
        if st.button("Send Query") and user_query:
            st.session_state.conversation.append({"role": "user", "content": user_query})

            # Generate the AI response (with spinner)
            if st.session_state.vector_store:
                with st.spinner("Generating response..."):
                    document_chain = create_stuff_documents_chain(llm, prompt)
                    retriever = st.session_state.vector_store.as_retriever()
                    retrieval_chain = create_retrieval_chain(retriever, document_chain)

                    # Query the vector store and get the relevant document chunks
                    st.session_state.response = retrieval_chain.invoke({'input': user_query})

                    # Simulate an AI response using LLM
                    ai_response = st.session_state.response['answer']
                    st.session_state.conversation.append({"role": "ai", "content": ai_response})

            else:
                st.error("The document vector store is not ready. Please process the documents first.")
        
        # Re-render chat history
        for message in st.session_state.conversation:
            display_message(message)

        # Display relevant chunks if response exists
        if st.session_state.response:
            with st.expander("Document Similarity Search"):
                for i, doc in enumerate(st.session_state.response["context"]):
                    st.write(doc.page_content)
                    st.write("--------------------------------")
    else:
        st.write("Upload and process documents on the left to start querying.")

# -----------------------------------
# Step 7: Display Chat Messages
# This function displays the user's query on the right and the AI's response on the left.
# It handles the layout and styling of the chat interface.
# -----------------------------------
def display_message(message):
    if message["role"] == "user":
        st.markdown(
            f"""
            <div style='display: flex; justify-content: flex-end; margin: 10px;'>
                <div style='background-color: #3B82F6; color: white; padding: 10px; border-radius: 10px; max-width: 60%; text-align: right;'>
                    {message['content']}
                </div>
                <img src="https://image.shutterstock.com/image-vector/man-icon-vector-260nw-1040084344.jpg" alt="User" width="40" height="40" style="margin-left: 10px;"/>
            </div>
            """, unsafe_allow_html=True)
    else:
        st.markdown(
            f"""
            <div style='display: flex; justify-content: flex-start; margin: 10px;'>
                <img src="https://image.shutterstock.com/image-vector/robot-icon-vector-260nw-1185116711.jpg" alt="AI" width="40" height="40" style="margin-right: 10px;"/>
                <div style='background-color: #555555; color: white; padding: 10px; border-radius: 10px; max-width: 60%; text-align: left;'>
                    {message['content']}
                </div>
            </div>
            """, unsafe_allow_html=True)

# -----------------------------------
# Main Execution Flow
# This is where the main execution happens. It calls the respective functions
# for document upload, processing, querying, and displaying chat.
# -----------------------------------
initialize_session_states()
left_col, right_col = st.columns([1, 2])

# Left column: Document upload and processing
with left_col:
    handle_document_upload()

# Right column: Chatbot UI and querying
with right_col:
    handle_query()
